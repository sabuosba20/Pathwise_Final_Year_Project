"""Safe, deterministic CV profile extraction for onboarding suggestions."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError


MAX_PDF_PAGES = 10
MAX_EXTRACTED_CHARACTERS = 100_000
MAX_DOCX_UNCOMPRESSED_BYTES = 12 * 1024 * 1024


class CvExtractionError(ValueError):
    """A user-safe CV validation or extraction error."""

    def __init__(self, message: str, status_code: int = 400):
        super().__init__(message)
        self.status_code = status_code


MAJOR_PATTERNS = [
    ("Data Science", (r"\bdata science\b", r"\bdata analytics?\b")),
    ("Software Engineering", (r"\bsoftware engineering\b",)),
    ("Information Technology", (r"\binformation technology\b", r"\bcomputing and it\b")),
    ("Cybersecurity", (r"\bcyber ?security\b", r"\binformation security\b")),
    ("Computer Science", (r"\bcomputer science\b", r"\bcomputing\b")),
    ("Business Management", (r"\bbusiness management\b", r"\bbusiness administration\b")),
    ("Finance", (r"\bfinance\b", r"\baccounting and finance\b")),
    ("Medicine", (r"\bmedicine\b", r"\bmedical science\b")),
    ("Psychology", (r"\bpsychology\b",)),
    ("Graphic Design", (r"\bgraphic design\b", r"\bvisual communication\b")),
    ("Marketing", (r"\bmarketing\b",)),
    ("Multimedia", (r"\bmultimedia\b", r"\bdigital media\b")),
    ("Linguistics", (r"\blinguistics\b",)),
    ("Education", (r"\beducation\b", r"\bteaching\b")),
    ("Engineering", (r"\bengineering\b",)),
]


SKILL_PATTERNS = [
    ("Python", (r"\bpython\b",)),
    ("Java", (r"\bjava\b",)),
    ("C++", (r"(?<!\w)c\+\+(?!\w)",)),
    ("C", (r"(?<![\w+#])c(?![\w+#])",)),
    ("JavaScript", (r"\bjavascript\b", r"\bjs\b")),
    ("TypeScript", (r"\btypescript\b",)),
    ("HTML", (r"\bhtml\b",)),
    ("CSS", (r"\bcss\b",)),
    ("Solidity", (r"\bsolidity\b",)),
    ("React", (r"\breact(?:\.js|js)?\b",)),
    ("Angular", (r"\bangular\b",)),
    ("Flask", (r"\bflask\b",)),
    ("Node.js", (r"\bnode(?:\.js|js)\b",)),
    ("Rust", (r"\brust\b",)),
    ("SQLite", (r"\bsqlite\b",)),
    ("SQL", (r"(?<!\w)sql(?!\w)",)),
    ("REST APIs", (r"\brest(?:ful)?\s+apis?\b",)),
    ("Git", (r"\bgit(?!hub)\b",)),
    ("GitHub", (r"\bgithub\b",)),
    ("Linux", (r"\blinux\b",)),
    ("Networking", (r"\bnetwork(?:ing|s)\b",)),
    ("Microsoft Excel", (r"\bmicrosoft excel\b", r"\bexcel\b")),
    ("Google Sheets", (r"\bgoogle sheets\b",)),
    ("Figma", (r"\bfigma\b",)),
    ("Machine Learning", (r"\bmachine learning\b",)),
    ("Data Analysis", (r"\bdata analy(?:sis|tics)\b",)),
    ("Object-Oriented Programming", (r"\bobject[- ]oriented programming\b", r"\boop\b")),
    ("Data Structures", (r"\bdata structures?\b",)),
    ("Algorithms", (r"\balgorithms?\b",)),
    ("Concurrent Programming", (r"\bconcurrent programming\b",)),
    ("Blockchain", (r"\bblockchain\b",)),
    ("Software Development Lifecycle", (r"\bsoftware development life ?cycle\b", r"\bsdlc\b")),
    ("Responsive UI Design", (r"\bresponsive (?:ui|user interface|web) design\b",)),
]


SECTION_HEADINGS = {
    "profile",
    "summary",
    "skills",
    "technical skills",
    "projects",
    "education",
    "experience",
    "work experience",
    "employment",
    "languages",
    "certifications",
    "awards",
    "references",
}


def _collapse_spaced_glyphs(line: str) -> str:
    """Repair PDFs that extract words as ``C o m p u t e r  S c i e n c e``."""
    tokens = re.findall(r"\S+", line)
    if len(tokens) < 4:
        return line
    single_character_ratio = sum(len(token) == 1 for token in tokens) / len(tokens)
    if single_character_ratio < 0.7:
        return line

    word_boundary = "\x01"
    repaired = re.sub(r"\s{2,}", word_boundary, line.strip())
    repaired = re.sub(r"\s+", "", repaired)
    return repaired.replace(word_boundary, " ")


def _normalized_lines(text: str) -> list[str]:
    lines = []
    for raw_line in text.replace("\x00", "").splitlines():
        repaired = _collapse_spaced_glyphs(raw_line)
        lines.append(" ".join(repaired.split()))
    return lines


def _section(text: str, wanted_headings: set[str]) -> str:
    collecting = False
    collected: list[str] = []
    for line in _normalized_lines(text):
        heading = line.strip(" :").casefold()
        if heading in SECTION_HEADINGS:
            if collecting:
                break
            collecting = heading in wanted_headings
            continue
        if collecting and line:
            collected.append(line)
    return " ".join(collected)


def _first_pattern_match(text: str, patterns):
    for label, expressions in patterns:
        if any(re.search(expression, text, flags=re.IGNORECASE) for expression in expressions):
            return label
    return None


def detect_profile_from_text(text: str) -> dict:
    """Return major and skills suggestions without retaining the source text."""
    cleaned_text = "\n".join(_normalized_lines(text))[:MAX_EXTRACTED_CHARACTERS]
    education_text = _section(cleaned_text, {"education"})
    skills_text = _section(cleaned_text, {"skills", "technical skills"})

    field_of_study = _first_pattern_match(education_text, MAJOR_PATTERNS)
    field_confidence = 0.96 if field_of_study else 0.0
    if not field_of_study:
        field_of_study = _first_pattern_match(cleaned_text, MAJOR_PATTERNS)
        field_confidence = 0.78 if field_of_study else 0.0

    skill_source = skills_text or cleaned_text
    skills = [
        label
        for label, expressions in SKILL_PATTERNS
        if any(re.search(expression, skill_source, flags=re.IGNORECASE) for expression in expressions)
    ]

    warnings = []
    if not field_of_study:
        warnings.append("No recognised field of study was found. Enter it manually.")
    if not skills:
        warnings.append("No recognised skills were found. Enter them manually.")

    return {
        "fieldOfStudy": field_of_study or "",
        "fieldConfidence": field_confidence,
        "skills": skills[:30],
        "warnings": warnings,
    }


def _extract_pdf_text(data: bytes) -> str:
    if not data.startswith(b"%PDF-"):
        raise CvExtractionError("The uploaded file is not a valid PDF.", 415)
    try:
        reader = PdfReader(BytesIO(data), strict=True)
        if reader.is_encrypted:
            raise CvExtractionError("Password-protected PDFs are not supported.")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise CvExtractionError(f"CVs may contain at most {MAX_PDF_PAGES} pages.")

        # Cap accumulated text as pages are read, before any regex normalisation
        # runs over it -- a page's compressed content stream can decode into a
        # far larger string than its on-disk size, so truncating only after
        # extraction would still pay the cost of processing the full blob.
        chunks = []
        remaining = MAX_EXTRACTED_CHARACTERS
        for page in reader.pages:
            if remaining <= 0:
                break
            page_text = (page.extract_text() or "")[:remaining]
            chunks.append(page_text)
            remaining -= len(page_text)
        return "\n".join(chunks)
    except CvExtractionError:
        raise
    except (PdfReadError, ValueError, TypeError) as error:
        raise CvExtractionError("The PDF could not be read safely.", 415) from error


def _extract_docx_text(data: bytes) -> str:
    if not data.startswith(b"PK"):
        raise CvExtractionError("The uploaded file is not a valid DOCX document.", 415)
    try:
        with ZipFile(BytesIO(data)) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise CvExtractionError("The uploaded file is not a valid DOCX document.", 415)
            if sum(item.file_size for item in archive.infolist()) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise CvExtractionError("The DOCX expands beyond the safe processing limit.")
        document = Document(BytesIO(data))
    except CvExtractionError:
        raise
    except (BadZipFile, ValueError, KeyError) as error:
        raise CvExtractionError("The DOCX document could not be read safely.", 415) from error

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def extract_cv_profile(stream, filename: str, maximum_bytes: int) -> dict:
    """Validate and parse an uploaded CV entirely in memory."""
    extension = Path(filename or "").suffix.casefold()
    if extension not in {".pdf", ".docx"}:
        raise CvExtractionError("Upload a PDF or DOCX CV.", 415)

    data = stream.read(maximum_bytes + 1)
    if not data:
        raise CvExtractionError("The uploaded CV is empty.")
    if len(data) > maximum_bytes:
        maximum_mb = maximum_bytes // (1024 * 1024)
        raise CvExtractionError(f"The CV must be {maximum_mb} MB or smaller.", 413)

    text = _extract_pdf_text(data) if extension == ".pdf" else _extract_docx_text(data)
    if len(text.strip()) < 80:
        raise CvExtractionError(
            "No readable text was found. Scanned or image-only CVs are not supported yet."
        )

    result = detect_profile_from_text(text)
    result["fileType"] = extension.removeprefix(".").upper()
    return result
